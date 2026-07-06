"""Генерация research-документа: отбор кейсов + синтез + сборка Word."""
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from loguru import logger

from db.base import get_session
from db.models import Trend, TrendCase
from llm.call_logger import llm_call_context
from research.llm_research import generate_research_synthesis, resolve_query_to_trends


@dataclass
class ResearchResult:
    query: str
    cases_used: int
    trends_covered: list[str]
    output_path: str


def _load_active_trends_list() -> list[dict]:
    with get_session() as s:
        trends = s.query(Trend).filter(Trend.status == "active").all()
        return [{"id": t.id, "name": t.name} for t in trends]


def _load_cases_for_trends(trend_ids: list[int], days: int | None, limit: int = 30) -> list[dict]:
    """Отбирает TrendCase по списку трендов. Исключает дубли и кейсы без тренда."""
    with get_session() as s:
        q = (
            s.query(TrendCase, Trend)
            .join(Trend, TrendCase.trend_id == Trend.id)
            .filter(TrendCase.trend_id.in_(trend_ids))
            .filter(TrendCase.is_duplicate == False)  # noqa: E712
        )
        if days:
            since = datetime.utcnow() - timedelta(days=days)
            q = q.filter(TrendCase.created_at >= since)

        q = q.order_by(TrendCase.created_at.desc()).limit(limit)

        result = []
        for tc, trend in q.all():
            result.append({
                "trend_name": trend.name,
                "company": tc.company,
                "case_title": tc.case_title,
                "description": tc.description,
                "how_it_works": tc.how_it_works,
                "value": tc.value,
                "market": tc.market,
            })
        return result


def _build_docx(
    query: str, synthesis: dict, cases: list[dict], trends_covered: list[str]
) -> Document:
    doc = Document()

    doc.add_heading(f"Исследование: {query}", level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Дата: {datetime.utcnow().strftime('%d.%m.%Y')}").italic = True
    meta.add_run(f"  •  Кейсов использовано: {len(cases)}").italic = True
    meta.add_run(f"  •  Тренды: {', '.join(trends_covered)}").italic = True

    doc.add_heading("Что происходит", level=2)
    doc.add_paragraph(synthesis["overview"])

    doc.add_heading("Технологии и подходы", level=2)
    doc.add_paragraph(synthesis["technology"])

    doc.add_heading("Кто и как реализует", level=2)
    doc.add_paragraph(synthesis["players"])

    doc.add_heading("Выводы", level=2)
    doc.add_paragraph(synthesis["conclusions"])

    doc.add_heading("Использованные кейсы", level=2)
    for c in cases:
        p = doc.add_paragraph()
        p.add_run(f"{c['case_title']}").bold = True
        p.add_run(f" — {c['company']} [{c['trend_name']}]")

    return doc


def generate_research(
    provider,
    query: str,
    days: int | None = None,
    output_path: str | None = None,
) -> "ResearchResult | None":
    """
    Генерирует research-документ по свободному текстовому запросу.
    Возвращает None если не найдено релевантных трендов или кейсов.
    """
    all_trends = _load_active_trends_list()

    logger.info(f"Research: resolving query '{query}' to trends...")
    with llm_call_context("resolve_query_to_trends", context_note=f"research: {query[:50]}"):
        trend_ids = resolve_query_to_trends(provider, query, all_trends)

    if not trend_ids:
        logger.warning(f"Research: no matching trends for query '{query}'")
        return None

    trend_names = [t["name"] for t in all_trends if t["id"] in trend_ids]
    logger.info(f"Research: matched trends: {trend_names}")

    cases = _load_cases_for_trends(trend_ids, days)
    if not cases:
        logger.warning(f"Research: no cases found for trends {trend_ids}")
        return None

    logger.info(f"Research: {len(cases)} cases found, generating synthesis...")
    with llm_call_context("generate_research_synthesis", context_note=f"research: {query[:50]}"):
        synthesis = generate_research_synthesis(provider, query, cases)

    if not output_path:
        Path("data/research").mkdir(parents=True, exist_ok=True)
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_query = "".join(c if c.isalnum() else "_" for c in query[:30])
        output_path = f"data/research/research_{safe_query}_{ts}.docx"

    doc = _build_docx(query, synthesis, cases, trend_names)
    doc.save(output_path)
    logger.info(f"Research saved: {output_path}")

    return ResearchResult(
        query=query,
        cases_used=len(cases),
        trends_covered=trend_names,
        output_path=output_path,
    )

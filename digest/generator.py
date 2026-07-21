from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from loguru import logger
from sqlalchemy import or_

from db.base import get_session
from db.models import NewsCard, Trend, TrendCase
from digest.llm_digest import generate_digest_analysis

COLOR_GREEN = RGBColor(0x1D, 0x9E, 0x75)
COLOR_DARK = RGBColor(0x2C, 0x2C, 0x2A)
COLOR_GRAY = RGBColor(0x88, 0x87, 0x80)
COLOR_LIGHT = RGBColor(0xB4, 0xB2, 0xA9)


@dataclass
class DigestResult:
    period_start: datetime
    period_end: datetime
    total_cases: int
    topics_count: int
    output_path: str


def generate_digest(
    provider,
    days: int = 7,
    output_path: str | None = None,
    max_cases: int = 50,
) -> DigestResult:
    period_end = datetime.utcnow()
    period_start = period_end - timedelta(days=days)

    if not output_path:
        Path("data/digests").mkdir(parents=True, exist_ok=True)
        ts = period_end.strftime("%Y%m%d")
        output_path = f"data/digests/digest_{ts}.docx"

    cases = _load_cases(period_start, max_cases)
    logger.info(f"Digest: {len(cases)} cases, period {period_start.date()}–{period_end.date()}")

    topics = _group_by_topic(cases)

    top_cases_for_context = sorted(
        cases, key=lambda c: c.get("relevance_score", 0), reverse=True
    )[:15]

    logger.info("LLM: единый аналитический вызов (главное + выводы по темам + векторы)...")
    analysis = generate_digest_analysis(provider, top_cases_for_context, topics) if cases else {
        "main_summary": "", "topic_conclusions": {}, "overall_conclusions": []
    }

    # Динамика — сравнение с прошлыми неделями (если они есть)
    past_snapshots = _load_past_snapshots(before=period_start, limit=3)

    current_index = []
    for topic, tcases in topics.items():
        for case in tcases:
            current_index.append({
                "company": case.get("company") or "—",
                "topic": topic,
                "title": case.get("case_title", "")[:100],
            })

    dynamics_points: list[str] = []
    if past_snapshots:
        logger.info(f"LLM: dynamics section (comparing with {len(past_snapshots)} past weeks)...")
        from digest.llm_digest import generate_dynamics_section
        dynamics_points = generate_dynamics_section(
            provider, current_index, analysis.get("overall_conclusions", []), past_snapshots
        )

    analysis["dynamics_points"] = dynamics_points

    doc = _build_docx(period_start, period_end, analysis, topics)
    doc.save(output_path)
    logger.info(f"Digest saved: {output_path}")

    _save_weekly_snapshot(period_start, period_end, analysis, topics)

    return DigestResult(
        period_start=period_start,
        period_end=period_end,
        total_cases=len(cases),
        topics_count=len(topics),
        output_path=output_path,
    )


def _load_cases(since: datetime, limit: int) -> list[dict]:
    with get_session() as session:
        rows = (
            session.query(TrendCase, NewsCard, Trend)
            .outerjoin(NewsCard, TrendCase.news_card_id == NewsCard.id)
            .outerjoin(Trend, TrendCase.trend_id == Trend.id)
            .filter(
                or_(
                    NewsCard.published_at >= since,
                    TrendCase.created_at >= since,
                )
            )
            .filter(
                or_(
                    NewsCard.relevance_score >= 85,
                    NewsCard.id == None,  # noqa: E711  — ручные кейсы
                )
            )
            .filter(TrendCase.is_duplicate == False)  # noqa: E712
            .order_by(NewsCard.relevance_score.desc().nullslast())
            .limit(limit)
            .all()
        )
        return [
            {
                "trend_name": tc.trend_name or (tr.name if tr else ""),
                "trend_category": tr.category if tr else "",
                "case_title": tc.case_title,
                "company": tc.company,
                "description": tc.description,
                "how_it_works": tc.how_it_works,
                "value": tc.value,
                "source_url": tc.source_url or (nc.post_url if nc else None),
                "source_title": nc.source_title if nc else "",
                "market": tc.market,
                "industry": tc.industry,
                "relevance_score": nc.relevance_score if nc else 0,
            }
            for tc, nc, tr in rows
        ]


def _save_weekly_snapshot(
    period_start: datetime,
    period_end: datetime,
    analysis: dict,
    topics: dict[str, list[dict]],
) -> None:
    """Сохраняет компактный снимок недели для будущего сравнения динамики."""
    import json
    from db.models import WeeklySnapshot

    compact_index = []
    for topic, cases in topics.items():
        for case in cases:
            compact_index.append({
                "company": case.get("company") or "—",
                "topic": topic,
                "title": case.get("case_title", "")[:100],
            })

    with get_session() as s:
        snapshot = WeeklySnapshot(
            period_start=period_start,
            period_end=period_end,
            main_summary=analysis.get("main_summary", ""),
            overall_conclusions=json.dumps(
                analysis.get("overall_conclusions", []), ensure_ascii=False
            ),
            compact_case_index=json.dumps(compact_index, ensure_ascii=False),
        )
        s.add(snapshot)
    logger.info(
        f"Weekly snapshot saved: {period_start.date()}–{period_end.date()}, "
        f"{len(compact_index)} cases indexed"
    )


def _load_past_snapshots(before: datetime, limit: int = 3) -> list[dict]:
    """Загружает последние N снимков ДО указанной даты (не включая текущий)."""
    import json
    from db.models import WeeklySnapshot

    with get_session() as s:
        snapshots = (
            s.query(WeeklySnapshot)
            .filter(WeeklySnapshot.period_end <= before)
            .order_by(WeeklySnapshot.period_end.desc())
            .limit(limit)
            .all()
        )

        result = []
        for snap in reversed(snapshots):  # от старых к новым
            try:
                case_index = json.loads(snap.compact_case_index or "[]")
                conclusions = json.loads(snap.overall_conclusions or "[]")
            except Exception:
                case_index, conclusions = [], []
            result.append({
                "period_label": (
                    f"{snap.period_start.strftime('%d.%m')}–"
                    f"{snap.period_end.strftime('%d.%m')}"
                ),
                "compact_case_index": case_index,
                "overall_conclusions": conclusions,
            })
        return result


def _group_by_topic(cases: list[dict]) -> dict[str, list[dict]]:
    seen: set[str] = set()
    groups: dict[str, list] = {}
    for c in cases:
        key = ((c.get("case_title") or "") + (c.get("company") or "")).lower()
        if key in seen:
            continue
        seen.add(key)
        topic = c.get("trend_category") or c.get("industry") or "Общее"
        groups.setdefault(topic, []).append(c)
    return dict(sorted(groups.items(), key=lambda x: -len(x[1])))


# ── Сборка документа ─────────────────────────────────────────────

def _build_docx(period_start, period_end, analysis: dict, topics: dict) -> Document:
    doc = Document()
    _setup_page(doc)
    _add_cover(doc, period_start, period_end, sum(len(v) for v in topics.values()), len(topics))

    # ── Главное за неделю ──────────────────────────────────────────
    _add_heading(doc, "Главное за неделю", 1)
    main_summary = analysis.get("main_summary", "")
    if main_summary:
        for paragraph_text in main_summary.split("\n\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                for run in p.runs:
                    run.font.size = Pt(10)
    else:
        doc.add_paragraph("Недостаточно данных для обзора.")

    # ── Динамика за месяц (если есть данные для сравнения) ─────────
    dynamics_points = analysis.get("dynamics_points", [])
    if dynamics_points:
        _add_heading(doc, "Динамика за месяц", 2)
        for point in dynamics_points:
            p = doc.add_paragraph(point)
            for run in p.runs:
                run.font.size = Pt(10)

    # ── Новости по темам ───────────────────────────────────────────
    _add_topics_section(doc, topics, analysis.get("topic_conclusions", {}))

    # ── Выводы и векторы изменений ─────────────────────────────────
    overall = analysis.get("overall_conclusions", [])
    if overall:
        _add_heading(doc, "Выводы и векторы изменений", 1)
        for point in overall:
            pb = doc.add_paragraph()
            pb.paragraph_format.left_indent = Cm(0.5)
            run = pb.add_run(f"• {point}")
            run.font.size = Pt(10)

    # ── Источники ───────────────────────────────────────────────────
    _add_sources_section(doc, topics)

    return doc


def _setup_page(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.5)
    s.top_margin = s.bottom_margin = Cm(2)


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Заголовок с явным форматированием — не зависит от встроенных стилей Word."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
    else:
        p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_GREEN
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_GREEN
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_DARK


def _add_cover(doc, start, end, n_cases, n_topics) -> None:
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Дайджест финтех-новостей")
    r.font.name = "Arial"; r.font.size = Pt(26)
    r.font.bold = True; r.font.color.rgb = COLOR_GREEN

    MONTHS_RU = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
        5: "мая", 6: "июня", 7: "июля", 8: "августа",
        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
    }
    period_str = (
        f"{start.day} {MONTHS_RU[start.month]} – "
        f"{end.day} {MONTHS_RU[end.month]} {end.year}"
    )
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(period_str)
    r2.font.name = "Arial"; r2.font.size = Pt(13); r2.font.color.rgb = COLOR_GRAY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"{n_cases} кейсов  ·  {n_topics} тем")
    r3.font.name = "Arial"; r3.font.size = Pt(10)
    r3.font.color.rgb = COLOR_LIGHT; r3.italic = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Подготовлено автоматически системой мониторинга Telegram-каналов")
    r4.font.name = "Arial"; r4.font.size = Pt(9)
    r4.font.color.rgb = COLOR_LIGHT; r4.italic = True

    _add_hline(doc)


def _add_topics_section(doc, topics: dict, topic_conclusions: dict) -> None:
    _add_heading(doc, "Новости по темам", 1)
    for topic, cases in topics.items():
        _add_heading(doc, topic, 2)
        conclusion = topic_conclusions.get(topic, "")
        if conclusion:
            pi = doc.add_paragraph()
            pi.paragraph_format.space_after = Pt(8)
            run = pi.add_run(conclusion)
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_DARK
            run.italic = True
        for c in cases:
            _add_case(doc, c)


def _add_case(doc, c: dict) -> None:
    pc = doc.add_paragraph()
    rt = pc.add_run(c.get("case_title") or "Без названия")
    rt.font.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = COLOR_DARK

    meta_parts = []
    if c.get("company"):
        meta_parts.append(c["company"])
    if c.get("market"):
        meta_parts.append(c["market"])
    if meta_parts:
        pm = doc.add_paragraph("  ·  ".join(meta_parts))
        pm.paragraph_format.left_indent = Cm(0.5)
        for run in pm.runs:
            run.font.size = Pt(9); run.font.color.rgb = COLOR_GRAY

    if c.get("description"):
        pd = doc.add_paragraph(c["description"])
        pd.paragraph_format.left_indent = Cm(0.5)
        for run in pd.runs:
            run.font.size = Pt(10)

    if c.get("value"):
        pv = doc.add_paragraph()
        pv.paragraph_format.left_indent = Cm(0.5)
        rv = pv.add_run("Ценность: ")
        rv.font.bold = True; rv.font.size = Pt(9); rv.font.color.rgb = COLOR_GREEN
        pv.add_run(c["value"]).font.size = Pt(9)

    if c.get("source_url"):
        pu = doc.add_paragraph()
        pu.paragraph_format.left_indent = Cm(0.5)
        ru = pu.add_run(c["source_url"])
        ru.font.size = Pt(8); ru.font.color.rgb = COLOR_LIGHT; ru.italic = True

    doc.add_paragraph()


def _add_sources_section(doc, topics) -> None:
    _add_heading(doc, "Источники", 1)
    seen: set[str] = set()
    i = 1
    for cases in topics.values():
        for c in cases:
            url = c.get("source_url", "")
            if url and url not in seen:
                seen.add(url)
                doc.add_paragraph(f"{i}. {url}")
                i += 1


def _add_hline(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1D9E75")
    pBdr.append(bottom)
    pPr.append(pBdr)
